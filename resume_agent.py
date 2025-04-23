from docx import Document
import pypandoc
from pathlib import Path


def resume_builder_agent(profile_data: dict, output_dir: str = "output") -> dict:
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)

    docx_path = output_path / "resume.docx"
    pdf_path = output_path / "resume.pdf"

    # 1. Build DOCX resume
    doc = Document()

    doc.add_heading(profile_data["name"], 0)
    doc.add_paragraph(f'{profile_data["email"]} | {profile_data["phone"]} | {profile_data["location"]}')

    doc.add_heading("Objective", level=1)
    doc.add_paragraph(profile_data["objective"])

    doc.add_heading("Education", level=1)
    for edu in profile_data["education"]:
        doc.add_paragraph(f'{edu["degree"]}, {edu["school"]} ({edu["dates"]})')

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(profile_data["skills"]))

    doc.add_heading("Experience", level=1)
    for job in profile_data["experience"]:
        doc.add_paragraph(f'{job["title"]} at {job["company"]} ({job["dates"]})')
        for bullet in job["description"].splitlines():
            doc.add_paragraph(f'• {bullet}', style='ListBullet')

    doc.add_heading("Projects", level=1)
    for proj in profile_data["projects"]:
        doc.add_paragraph(f'{proj["title"]}: {proj["description"]}')

    doc.add_heading("Certifications", level=1)
    for cert in profile_data["certifications"]:
        doc.add_paragraph(f'• {cert}', style='ListBullet')

    doc.save(docx_path)

    # 2. Convert DOCX to PDF using Pandoc (must be installed)
    try:
        pypandoc.convert_file(str(docx_path), "pdf", outputfile=str(pdf_path))
    except Exception as e:
        return {
            "error": f"PDF conversion failed: {str(e)}",
            "docx_file": str(docx_path),
            "pdf_file": None
        }

    return {
        "docx_file": str(docx_path),
        "pdf_file": str(pdf_path)
    }
